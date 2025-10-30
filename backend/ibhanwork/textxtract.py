# ==============================================================
# 📦  Install required packages (run once)
# ==============================================================
# !pip install -q PyPDF2==3.0.1 pdfplumber==0.10.1 python-docx==1.1.0 tqdm==4.66.2

# Try to install antiword for .doc support (optional)
# If this fails we will simply skip .doc files.
# !apt-get -qq update && apt-get -qq install -y antiword


from google.colab import files
uploaded = files.upload()          # choose your zip file

if not uploaded:
    raise SystemExit(" No file uploaded – stop.")
zip_path = list(uploaded.keys())[0]   # e.g. "my_resumes.zip"
print(f" Uploaded: {zip_path}")


import os, zipfile, shutil, sys, subprocess
from pathlib import Path
import pdfplumber, PyPDF2
from docx import Document
from tqdm import tqdm
from datetime import datetime
import IPython.display as ipd

def extract_text_from_pdf(fp):
    """pdfplumber first (better layout), fallback to PyPDF2."""
    txt = ""
    try:
        with pdfplumber.open(fp) as pdf:
            for page in pdf.pages:
                page_txt = page.extract_text()
                if page_txt:
                    txt += page_txt + "\n"
    except Exception:
        pass          # pdfplumber failed – try PyPDF2
    if not txt.strip():
        try:
            with open(fp, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_txt = page.extract_text()
                    if page_txt:
                        txt += page_txt + "\n"
        except Exception as e:
            print(f"❌ PDF error ({fp}): {e}")
    return txt.strip()

def extract_text_from_docx(fp):
    """Extract text from a .docx file."""
    try:
        doc = Document(fp)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f" DOCX error ({fp}): {e}")
        return ""

def extract_text_from_doc(fp):
    """
    Extract text from a legacy .doc file using the `antiword` CLI tool.
    If `antiword` is not available, we return a friendly message.
    """
    try:
        # `antiword -w 0` returns plain UTF‑8 text
        result = subprocess.run(
            ["antiword", "-w", "0", fp],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True,
        )
        return result.stdout.decode("utf-8", errors="ignore")
    except FileNotFoundError:
        # antiword not installed
        return ("  antiword not found – .doc files cannot be processed. "
                "Please install antiword or convert the files to .docx/.pdf.")
    except subprocess.CalledProcessError as e:
        return (f"  antiword failed on {fp}: {e.stderr.decode('utf-8', errors='ignore')}"
                "\nPlease check the file format.")
    except Exception:
        return ("  Unexpected error while processing .doc file. "
                "File will be skipped.")

def clean_temp_files():
    """Remove the temporary folder used for extracting files from the zip."""
    try:
        if os.path.exists("temp"):
            shutil.rmtree("temp")
    except Exception:
        pass

def safe_name(name):
    """Return a filename that contains only alphanumerics, space, dash or underscore."""
    cleaned = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).rstrip()
    return cleaned if cleaned else None

def download_file(path):
    """Trigger a download of a single file."""
    files.download(str(path))


extract_root = Path("/content/unzipped")
extract_root.mkdir(parents=True, exist_ok=True)

with zipfile.ZipFile(zip_path, "r") as zip_ref:
    zip_ref.extractall(extract_root)


d_folder = None
for p in extract_root.rglob("*"):
    if p.is_dir() and p.name.lower() == "d":
        d_folder = p
        break

if d_folder is None:
    raise SystemExit(" Could not find a folder named 'd' inside the zip.")
print(f" Found folder 'd' at: {d_folder}")


supported_exts = {".pdf", ".docx", ".doc"}
resume_paths = [p for p in d_folder.rglob("*")
                if p.is_file() and p.suffix.lower() in supported_exts]

if not resume_paths:
    raise SystemExit(" No .pdf/.docx/.doc files found inside folder 'd'.")

print(f"📋 {len(resume_paths)} resume files will be processed.")


output_dir = Path("/content/extracted_texts")
output_dir.mkdir(parents=True, exist_ok=True)

extracted = []         
errors = []

for i, src_path in enumerate(tqdm(resume_paths, desc="Extracting"), start=1):
    ext = src_path.suffix.lower()
    try:
        if ext == ".pdf":
            txt = extract_text_from_pdf(str(src_path))
        elif ext == ".docx":
            txt = extract_text_from_docx(str(src_path))
        else:   # .doc
            txt = extract_text_from_doc(str(src_path))
    except Exception as e:
        errors.append(f"{src_path}: {e}")
        continue

    
    base = src_path.stem
    safe = safe_name(base) or f"resume_{i}"
    out_file = output_dir / f"{safe}.txt"

    header = (f"=== EXTRACTED FROM: {src_path.relative_to(d_folder)} ===\n"
              f"=== DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ===\n"
              f"=== SOURCE: folder 'd' inside zip ===\n\n")
    try:
        out_file.write_text(header + txt, encoding="utf-8")
        extracted.append((src_path.name, out_file.name))
    except Exception as e:
        errors.append(f"Write error for {out_file}: {e}")

# Clean up the temporary folder we created for the zip extraction
clean_temp_files()


print("\n" + "="*60)
print(" EXTRACTION SUMMARY")
print("="*60)

if extracted:
    print(f" Created {len(extracted)} .txt files in `{output_dir}`")
    print("\n  Sample of created files:")
    for orig, txt_name in extracted[:10]:
        print(f"   • {orig} → {txt_name}")
    if len(extracted) > 10:
        print(f"   ... and {len(extracted)-10} more")
else:
    print(" No files were written.")

if errors:
    print(f"\n  {len(errors)} error(s) occurred (showing up to 5):")
    for e in errors[:5]:
        print(f"   • {e}")
    if len(errors) > 5:
        print(f"   ... and {len(errors)-5} more")

# Show a short preview of the first resume (optional)
if extracted:
    first_txt = output_dir / extracted[0][1]
    print("\n Preview of the first extracted resume:")
    try:
        preview = first_txt.read_text(encoding="utf-8")[:1000]
        print("-"*60)
        print(preview + ("…" if len(preview) == 1000 else ""))
        print("-"*60)
    except Exception as e:
        print(f"Could not read preview: {e}")

print("\n Click a button to download any of the extracted .txt files:")
for txt_name in sorted(os.listdir(output_dir)):
    path = output_dir / txt_name
    button_html = f'''
    <form action="/download" method="post" target="_blank">
        <input type="hidden" name="filename" value="{path}">
        <input type="submit" value="Download {txt_name}" style="margin:2px;">
    </form>
    '''
    ipd.display(ipd.HTML(button_html))

print("\n All done! The folder `extracted_texts/` now contains one .txt per resume.")
